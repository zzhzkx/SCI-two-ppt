"""论文解析器 - 支持PDF和Word文档."""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Optional
import json
import re


async def parse_papers(papers: list[str]) -> dict:
    """解析论文（支持PDF和Word文档）。

    Args:
        papers: 文件路径列表（支持 .pdf 和 .docx）

    Returns:
        dict: 解析结果和质量报告
    """
    results = []
    errors = []

    for paper_path in papers:
        p = Path(paper_path)

        if not p.exists():
            errors.append(f"文件不存在: {paper_path}")
            continue

        try:
            if p.suffix.lower() == ".pdf":
                result = await _parse_pdf(p)
            elif p.suffix.lower() in [".docx", ".doc"]:
                result = await _parse_docx(p)
            else:
                errors.append(f"不支持的文件格式: {p.suffix}")
                continue
            results.append(result)
        except Exception as e:
            errors.append(f"解析失败 {paper_path}: {str(e)}")

    quality_report = _generate_quality_report(results, errors)
    return {"papers": results, "quality_report": quality_report}


async def _parse_single_paper(pdf_path: Path) -> dict:
    """解析单篇论文."""
    doc = fitz.open(str(pdf_path))

    # 提取全文
    full_text = ""
    for page in doc:
        full_text += page.get_text()

    # 提取各部分
    title = _extract_title(full_text, pdf_path)
    abstract = _extract_section(full_text, ["abstract", "summary"])
    methods = _extract_section(full_text, ["method", "methodology", "approach", "experiment"])
    results = _extract_section(full_text, ["result", "finding", "outcome", "analysis"])

    # 提取关键发现和创新点
    key_findings = _extract_key_findings(full_text)
    innovations = _extract_innovations(full_text)

    doc.close()

    return {
        "path": str(pdf_path),
        "title": title,
        "abstract": abstract[:2000],
        "methods": methods[:2000],
        "results": results[:2000],
        "figures": [],
        "key_findings": key_findings,
        "innovations": innovations
    }


async def _parse_docx(docx_path: Path) -> dict:
    """解析Word文档."""
    doc = Document(str(docx_path))

    # 提取全文
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # 提取各部分
    title = _extract_title_from_docx(doc, docx_path)

    # 支持中英文关键词
    abstract = _extract_section(full_text, ["abstract", "摘要"])
    methods = _extract_section(full_text, ["method", "方法", "实验", "实验方法", "二、", "2.", "1.1", "1.2"])
    results = _extract_section(full_text, ["result", "结果", "实验结果", "结论", "三、", "3."])

    # 提取关键发现和创新点
    key_findings = _extract_key_findings(full_text)
    innovations = _extract_innovations(full_text)

    return {
        "path": str(docx_path),
        "title": title,
        "abstract": abstract[:2000],
        "methods": methods[:2000],
        "results": results[:2000],
        "figures": [],
        "key_findings": key_findings,
        "innovations": innovations
    }


def _extract_title_from_docx(doc: Document, docx_path: Path) -> str:
    """从Word文档提取标题."""
    # 尝试从标题样式提取
    for para in doc.paragraphs[:20]:
        if para.style and para.style.name and 'Heading' in para.style.name:
            if para.text.strip() and len(para.text.strip()) > 10:
                return para.text.strip()

    # 尝试从第一段非空内容提取
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        # 跳过空行和太短的行
        if len(text) > 15 and len(text) < 200 and not text.startswith(('学生', '学号', '专业')):
            return text

    # 返回文件名
    return docx_path.stem.replace('-', ' ').replace('_', ' ')


def _extract_title(text: str, pdf_path: Path) -> str:
    """提取论文标题."""
    # 尝试从文本开头提取
    lines = text.strip().split('\n')

    # 通常标题是前几行中最长的非空行
    for line in lines[:10]:
        line = line.strip()
        if len(line) > 20 and not line.startswith(('http', 'www', 'Abstract', 'abstract')):
            return line

    # 如果找不到，返回文件名
    return pdf_path.stem.replace('_', ' ').replace('-', ' ')


def _extract_section(text: str, keywords: list[str]) -> str:
    """提取指定章节内容."""
    text_lower = text.lower()
    start_idx = -1

    # 查找章节开始位置
    for keyword in keywords:
        # 尝试匹配 "Abstract" 或 "Abstract:" 等格式
        patterns = [
            rf'\b{keyword}\b\s*[:\.]',
            rf'\b{keyword}\b\s*\n',
        ]
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                start_idx = match.start()
                break
        if start_idx != -1:
            break

    if start_idx == -1:
        return ""

    # 提取到下一个主要章节或文末
    section_text = text[start_idx:]
    next_section_patterns = [
        r'\b(?:introduction|method|result|discussion|conclusion|reference|acknowledgment)\b',
    ]

    for pattern in next_section_patterns:
        match = re.search(pattern, section_text[100:], re.IGNORECASE)
        if match:
            section_text = section_text[:100 + match.start()]

    return section_text[:3000].strip()


def _extract_key_findings(text: str) -> list[str]:
    """提取关键发现."""
    findings = []

    # 查找包含关键发现的句子
    patterns = [
        r'(?:we found|our results show|demonstrated|revealed|indicated|suggested)\s+[^.]+\.',
        r'(?:the results|findings|analysis)\s+(?:show|indicate|suggest|demonstrate)\s+[^.]+\.',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches[:5]:  # 每个模式最多5个
            if len(match) > 30 and len(match) < 300:
                findings.append(match.strip())

    return findings[:10]  # 最多10个关键发现


def _extract_innovations(text: str) -> list[str]:
    """提取创新点."""
    innovations = []

    # 查找包含创新性的句子
    patterns = [
        r'(?:novel|new|innovative|proposed|developed|introduced)\s+[^.]+\.',
        r'(?:for the first time|first study|first report)\s+[^.]+\.',
        r'(?:our contribution|we contribute|main contribution)\s+[^.]+\.',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches[:3]:
            if len(match) > 30 and len(match) < 300:
                innovations.append(match.strip())

    return innovations[:5]  # 最多5个创新点


def _generate_quality_report(papers: list[dict], errors: list[str]) -> str:
    """生成质量报告."""
    report = "# 论文解析质量报告\n\n"

    if errors:
        report += "## 错误\n"
        for error in errors:
            report += f"- {error}\n"
        report += "\n"

    if papers:
        report += "## 解析结果\n"
        for i, paper in enumerate(papers, 1):
            report += f"\n### 论文 {i}: {paper.get('title', '未知')}\n"
            report += f"- 路径: {paper.get('path', 'N/A')}\n"
            report += f"- 摘要长度: {len(paper.get('abstract', ''))} 字符\n"
            report += f"- 方法长度: {len(paper.get('methods', ''))} 字符\n"
            report += f"- 结果长度: {len(paper.get('results', ''))} 字符\n"
            report += f"- 关键发现: {len(paper.get('key_findings', []))} 条\n"
            report += f"- 创新点: {len(paper.get('innovations', []))} 条\n"

            # 质量评估
            quality = _assess_quality(paper)
            report += f"- 质量评估: {quality}\n"

    report += "\n## 总结\n"
    report += f"- 成功解析: {len(papers)} 篇\n"
    report += f"- 失败: {len(errors)} 篇\n"

    if len(papers) == 0:
        report += "\n**警告**: 未成功解析任何论文\n"

    return report


def _assess_quality(paper: dict) -> str:
    """评估单篇论文解析质量."""
    issues = []

    if not paper.get('abstract'):
        issues.append("缺少摘要")
    if not paper.get('methods'):
        issues.append("缺少方法")
    if not paper.get('results'):
        issues.append("缺少结果")
    if len(paper.get('key_findings', [])) == 0:
        issues.append("未提取到关键发现")
    if len(paper.get('innovations', [])) == 0:
        issues.append("未提取到创新点")

    if not issues:
        return "良好"
    elif len(issues) <= 2:
        return "一般 - " + ", ".join(issues)
    else:
        return "较差 - " + ", ".join(issues)
